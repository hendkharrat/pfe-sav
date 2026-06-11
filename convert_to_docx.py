"""
Convert rapport_pfe_corrige.md to rapport_pfe_corrige.docx
Handles: headings (H1-H4), tables, bullet lists, numbered lists,
         bold/inline-code inline formatting, hr separators, paragraphs.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = r"C:\Users\medam\Desktop\PFE Hend\sav-manager-frontend-setup\rapport_pfe_corrige.md"
OUT_PATH = r"C:\Users\medam\Desktop\PFE Hend\sav-manager-frontend-setup\rapport_pfe_corrige.docx"

# ---------- helpers ----------

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_inline_runs(para, text):
    """
    Parse **bold**, `code`, and plain text, adding runs accordingly.
    """
    # Pattern: **bold** or `code` or plain segments
    token_re = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`)')
    pos = 0
    for m in token_re.finditer(text):
        # plain text before this token
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.size = Pt(11)
        if m.group(0).startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        else:  # backtick code
            run = para.add_run(m.group(3))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        pos = m.end()
    # remaining plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(11)

def style_heading(para, level):
    sizes  = {1: 20, 2: 16, 3: 14, 4: 12}
    colors = {
        1: RGBColor(0x1F, 0x39, 0x7A),
        2: RGBColor(0x2E, 0x74, 0xB5),
        3: RGBColor(0x1F, 0x60, 0x7A),
        4: RGBColor(0x40, 0x40, 0x40),
    }
    for run in para.runs:
        run.font.size  = Pt(sizes.get(level, 12))
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
        run.font.bold  = True
    para.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    para.paragraph_format.space_after  = Pt(4)

# ---------- table detection ----------

def is_table_row(line):
    return line.strip().startswith('|') and line.strip().endswith('|')

def is_separator_row(line):
    return is_table_row(line) and re.match(r'^\|[-| :]+\|$', line.strip())

def parse_table_row(line):
    parts = line.strip().strip('|').split('|')
    return [p.strip() for p in parts]

# ---------- list helpers ----------

def is_bullet(line):
    return re.match(r'^(\s*)[-*] ', line)

def is_numbered(line):
    return re.match(r'^(\s*)\d+\. ', line)

def bullet_indent(line):
    m = re.match(r'^(\s*)[-*] ', line)
    return len(m.group(1)) // 2 if m else 0

# ---------- paragraph block parser ----------

def collect_table_block(lines, start):
    """Return (rows_list, next_index). rows_list = list of cell-string-lists."""
    rows   = []
    i      = start
    header = None
    while i < len(lines):
        line = lines[i]
        if is_separator_row(line):
            i += 1
            continue
        if is_table_row(line):
            row = parse_table_row(line)
            rows.append(row)
            i += 1
        else:
            break
    return rows, i

# ---------- main render ----------

def render(doc, lines):
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # horizontal rule ---
        if re.match(r'^-{3,}$', stripped):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # headings
        hm = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2)
            style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}.get(level, 'Heading 4')
            para  = doc.add_paragraph(style=style)
            add_inline_runs(para, text)
            style_heading(para, level)
            i += 1
            continue

        # table block
        if is_table_row(stripped):
            rows, i = collect_table_block(lines, i)
            if not rows:
                continue
            # determine col count from max row length
            col_count = max(len(r) for r in rows)
            # normalise row lengths
            rows = [r + [''] * (col_count - len(r)) for r in rows]
            tbl = doc.add_table(rows=len(rows), cols=col_count)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    para = cell.paragraphs[0]
                    para.clear()
                    add_inline_runs(para, cell_text)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)
                    set_cell_borders(cell)
                    if ri == 0:
                        set_cell_bg(cell, 'D6E4F0')
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                    else:
                        for run in para.runs:
                            run.font.size = Pt(10)
            doc.add_paragraph()  # spacing after table
            continue

        # bullet list item
        if is_bullet(line):
            indent = bullet_indent(line)
            text   = re.sub(r'^\s*[-*] ', '', line)
            level_style = 'List Bullet 2' if indent > 0 else 'List Bullet'
            try:
                para = doc.add_paragraph(style=level_style)
            except Exception:
                para = doc.add_paragraph(style='List Bullet')
            add_inline_runs(para, text.strip())
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # numbered list item
        if is_numbered(line):
            text = re.sub(r'^\s*\d+\. ', '', line)
            try:
                para = doc.add_paragraph(style='List Number')
            except Exception:
                para = doc.add_paragraph(style='Normal')
            add_inline_runs(para, text.strip())
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # diagram placeholder  [Insérer ici ...]
        if stripped.startswith('[Insérer'):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.italic = True
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            # light grey background via shading on the paragraph
            pPr  = para._p.get_or_add_pPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'F2F2F2')
            pPr.append(shd)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        # plain paragraph — collect continuation lines
        text_parts = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if (not nxt or
                    re.match(r'^#{1,4}\s', nxt) or
                    is_table_row(nxt) or
                    is_bullet(lines[j]) or
                    is_numbered(lines[j]) or
                    re.match(r'^-{3,}$', nxt) or
                    nxt.startswith('[Insérer')):
                break
            text_parts.append(nxt)
            j += 1
        full_text = ' '.join(text_parts)
        para = doc.add_paragraph(style='Normal')
        add_inline_runs(para, full_text)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        i = j

# ---------- setup & run ----------

def setup_styles(doc):
    """Ensure heading styles exist and set base Normal font."""
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

def main():
    with open(MD_PATH, encoding='utf-8') as f:
        lines = f.readlines()
    lines = [l.rstrip('\n') for l in lines]

    doc = Document()
    setup_styles(doc)
    render(doc, lines)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")

if __name__ == '__main__':
    main()
